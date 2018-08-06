from tkinter import *
from tkinter import ttk
import pandas as pd
import numpy as np
import Pmw
import csv
import openpyxl
import os
import datetime
from datetime import date
from datetime import time
from datetime import datetime,timedelta
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import asksaveasfilename
import docx
import sys
import comtypes.client
import smtplib
import imapclient
import re
import email
import mimetypes
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email import encoders
import shutil
from shutil import copyfile
import glob
import locale

tab1 = Tk()
#nb = ttk.Notebook(root)
tab1.title('DiRA')
#tab1 = Frame(nb)
#tab2 = Frame(nb)
#tab3 = Frame(nb)
#tab4 = Frame(nb)
#tab5 = Frame(nb)
#tab6 = Frame(nb)

#AutoSend NOTA
#tab1
a=StringVar()
b=StringVar()
c=StringVar()
d=StringVar()
e=StringVar()
f=StringVar()
g=StringVar()
h=StringVar()
i=StringVar()
j=StringVar()
k=StringVar()
l=StringVar()
m=StringVar()
n=StringVar()
o=StringVar()
p=StringVar()
q=StringVar()
r=StringVar()
s=StringVar()
t=StringVar()
u=StringVar()
v=StringVar()
w=StringVar()
x=StringVar()
y=StringVar()
z=StringVar()
aa=StringVar()
ab=StringVar()
ac=StringVar()
ad=StringVar()
ae=StringVar()
af=StringVar()
ag=StringVar()
ah=StringVar()
ai=StringVar()
aj=StringVar()
ak=StringVar()
al=StringVar()
am=StringVar()
an=StringVar()
ao=StringVar()
ap=StringVar()
aq=StringVar()
ar=StringVar()
at=StringVar()
au=StringVar()
av=StringVar()
aw=StringVar()
ax=StringVar()
ay=StringVar()
az=StringVar()

Label(tab1, text="No Nota:").grid(row=1,column=0)
notaEntry=Entry(tab1, width=30, bg='white',textvariable=a)
notaEntry.grid(row=1,column=1)
Label(tab1, text="Perihal:").grid(row=2,column=0)

opt_perihal = Pmw.OptionMenu(tab1, menubutton_textvariable=b,items=('L15','Metric Ton','Reservasi','ShipTo bentrok dengan Plumpang'),menubutton_width=16)
opt_perihal.grid(row=2,column=1)

#Label Tabel
noLabel=Label(tab1, text="NO")
nopoLabel=Label(tab1, text="NOPOL")
noloLabel=Label(tab1, text="LO")
produkLabel=Label(tab1, text="PRODUK")
mlLabel = Label(tab1,text='*MT|L15')
voLabel=Label(tab1, text="VOLUME(L)")
ketLabel=Label(tab1, text="KETERANGAN")
noLabel.grid(row=3,column=0)
nopoLabel.grid(row=3,column=1)
noloLabel.grid(row=3,column=2)
produkLabel.grid(row=3,column=3)
mlLabel.grid(row=3,column=4)
voLabel.grid(row=3,column=5)
ketLabel.grid(row=3,column=6)

#Entry Tabel
noEntry1 = Entry(tab1, width = 5,bg='white',textvariable=c)
noEntry2 = Entry(tab1, width = 5,bg='white',textvariable=d)
noEntry3 = Entry(tab1, width = 5,bg='white',textvariable=e)
noEntry4 = Entry(tab1, width = 5,bg='white',textvariable=f)
noEntry5 = Entry(tab1, width = 5,bg='white',textvariable=g)
noEntry6 = Entry(tab1, width = 5,bg='white',textvariable=h)
noEntry7 = Entry(tab1, width = 5,bg='white',textvariable=i)
nopolEntry1 = Entry(tab1, width=20,bg='white',textvariable=j)
nopolEntry2 = Entry(tab1, width=20,bg='white',textvariable=k)
nopolEntry3 = Entry(tab1, width=20,bg='white',textvariable=l)
nopolEntry4 = Entry(tab1, width=20,bg='white',textvariable=m)
nopolEntry5 = Entry(tab1, width=20,bg='white',textvariable=n)
nopolEntry6 = Entry(tab1, width=20,bg='white',textvariable=o)
nopolEntry7 = Entry(tab1, width=20,bg='white',textvariable=p)
noloEntry1 = Entry(tab1, width=20,bg='white',textvariable=q)
noloEntry2 = Entry(tab1, width=20,bg='white',textvariable=r)
noloEntry3 = Entry(tab1, width=20,bg='white',textvariable=s)
noloEntry4 = Entry(tab1, width=20,bg='white',textvariable=t)
noloEntry5 = Entry(tab1, width=20,bg='white',textvariable=u)
noloEntry6 = Entry(tab1, width=20,bg='white',textvariable=v)
noloEntry7 = Entry(tab1, width=20,bg='white',textvariable=w)
prodEntry1 = Entry(tab1, width=20,bg='white',textvariable=x)
prodEntry2 = Entry(tab1, width=20,bg='white',textvariable=y)
prodEntry3 = Entry(tab1, width=20,bg='white',textvariable=z)
prodEntry4 = Entry(tab1, width=20,bg='white',textvariable=aa)
prodEntry5 = Entry(tab1, width=20,bg='white',textvariable=ab)
prodEntry6 = Entry(tab1, width=20,bg='white',textvariable=ac)
prodEntry7 = Entry(tab1, width=20,bg='white',textvariable=ad)
volEntry1 = Entry(tab1,width=10,bg='white',textvariable=ae)
volEntry2 = Entry(tab1,width=10,bg='white',textvariable=af)
volEntry3 = Entry(tab1,width=10,bg='white',textvariable=ag)
volEntry4 = Entry(tab1,width=10,bg='white',textvariable=ah)
volEntry5 = Entry(tab1,width=10,bg='white',textvariable=ai)
volEntry6 = Entry(tab1,width=10,bg='white',textvariable=aj)
volEntry7 = Entry(tab1,width=10,bg='white',textvariable=ak)
ketEntry=Entry(tab1,width=30,bg='white',textvariable=al)
ketEntry2=Entry(tab1,width=30,bg='white',textvariable=au)
ketEntry3=Entry(tab1,width=30,bg='white',textvariable=av)
ketEntry4=Entry(tab1,width=30,bg='white',textvariable=aw)
ketEntry5=Entry(tab1,width=30,bg='white',textvariable=ax)
ketEntry6=Entry(tab1,width=30,bg='white',textvariable=ay)
ketEntry7=Entry(tab1,width=30,bg='white',textvariable=az)
mlEntry1=Entry(tab1,width=10,bg='white',textvariable=am)
mlEntry2=Entry(tab1,width=10,bg='white',textvariable=an)
mlEntry3=Entry(tab1,width=10,bg='white',textvariable=ao)
mlEntry4=Entry(tab1,width=10,bg='white',textvariable=ap)
mlEntry5=Entry(tab1,width=10,bg='white',textvariable=aq)
mlEntry6=Entry(tab1,width=10,bg='white',textvariable=ar)
mlEntry7=Entry(tab1,width=10,bg='white',textvariable=at)
#row4
noEntry1.grid(row=4)
nopolEntry1.grid(row=4,column=1)
noloEntry1.grid(row=4,column=2)
prodEntry1.grid(row=4,column=3)
mlEntry1.grid(row=4,column=4)
volEntry1.grid(row=4,column=5)
ketEntry.grid(row=4,column=6)
#row5
noEntry2.grid(row=5)
nopolEntry2.grid(row=5,column=1)
noloEntry2.grid(row=5,column=2)
prodEntry2.grid(row=5,column=3)
mlEntry2.grid(row=5,column=4)
volEntry2.grid(row=5,column=5)
ketEntry2.grid(row=5,column=6)
#row6
noEntry3.grid(row=6)
nopolEntry3.grid(row=6,column=1)
noloEntry3.grid(row=6,column=2)
prodEntry3.grid(row=6,column=3)
mlEntry3.grid(row=6,column=4)
volEntry3.grid(row=6,column=5)
ketEntry3.grid(row=6,column=6)
#row7
noEntry4.grid(row=7)
nopolEntry4.grid(row=7,column=1)
noloEntry4.grid(row=7,column=2)
prodEntry4.grid(row=7,column=3)
mlEntry4.grid(row=7,column=4)
volEntry4.grid(row=7,column=5)
ketEntry4.grid(row=7,column=6)
#row8
noEntry5.grid(row=8)
nopolEntry5.grid(row=8,column=1)
noloEntry5.grid(row=8,column=2)
prodEntry5.grid(row=8,column=3)
mlEntry5.grid(row=8,column=4)
volEntry5.grid(row=8,column=5)
ketEntry5.grid(row=8,column=6)
#row9
noEntry6.grid(row=9)
nopolEntry6.grid(row=9,column=1)
noloEntry6.grid(row=9,column=2)
prodEntry6.grid(row=9,column=3)
mlEntry6.grid(row=9,column=4)
volEntry6.grid(row=9,column=5)
ketEntry6.grid(row=9,column=6)
#row10
noEntry7.grid(row=10)
nopolEntry7.grid(row=10,column=1)
noloEntry7.grid(row=10,column=2)
prodEntry7.grid(row=10,column=3)
mlEntry7.grid(row=10,column=4)
volEntry7.grid(row=10,column=5)
ketEntry7.grid(row=10,column=6)
def runShit():
	today = date.today()

	#1.interpolasi nomor nota
	thisMonth = today.strftime('%m')
	thisYear = today.strftime('%Y')

	def rome(n):
		if n == '01':
			rom = 'I'
			return rom
		if n == '02':
			rom = 'II'
			return rom
		if n == '03':
			rom = 'III'
			return rom
		if n == '04':
			rom = 'IV'
			return rom
		if n == '05':
			rom = 'V'
			return rom
		if n == '06':
			rom = 'VI'
			return rom
		if n == '07':
			rom = 'VII'
			return rom
		if n == '08':
			rom = 'VIII'
			return rom
		if n == '09':
			rom = 'IX'
			return rom
		if n == '10':
			rom = 'X'
			return rom
		if n == '11':
			rom = 'XI'
			return rom
		if n == '12':
			rom = 'XII'
			return rom

	romanizing = rome(thisMonth)
	notaYear = today.strftime('%Y')
	nomerNota = a.get() #dimanipulasi melalui email kedepannya

	noNota = ' {0}/PLP-MNJ/{1}/{2}'.format(nomerNota,romanizing,notaYear)

	#2.interpolasi Perihal (conditional)
	perihal = b.get()
	#3.interpolasi Tabel
	nomer1 = c.get()
	nomer2 = d.get()
	nomer3 = e.get()
	nomer4 = f.get()
	nomer5 = g.get()
	nomer6 = h.get()
	nomer7 = i.get()
	mobil1 = j.get()
	mobil2 = k.get()
	mobil3 = l.get()
	mobil4 = m.get()
	mobil5 = n.get()
	mobil6 = o.get()
	mobil7 = p.get()
	nomerLO1 = q.get()
	nomerLO2 = r.get()
	nomerLO3 = s.get()
	nomerLO4 = t.get()
	nomerLO5 = u.get()
	nomerLO6 = v.get()
	nomerLO7 = w.get()
	produkLO1 = x.get()
	produkLO2 = y.get()
	produkLO3 = z.get()
	produkLO4 = aa.get()
	produkLO5 = ab.get()
	produkLO6 = ac.get()
	produkLO7 = ad.get()
	ml1=am.get()
	ml2=an.get()
	ml3=ao.get()
	ml4=ap.get()
	ml5=aq.get()
	ml6=ar.get()
	ml7=at.get()
	volume1 = ae.get()
	volume2 = af.get()
	volume3 = ag.get()
	volume4 = ah.get()
	volume5 = ai.get()
	volume6 = aj.get()
	volume7 = ak.get()
	keterangan= al.get()
	keterangan2= au.get()
	keterangan3= av.get()
	keterangan4= aw.get()
	keterangan5= ax.get()
	keterangan6= ay.get()
	keterangan7= az.get()
	#2.1 L15
	#L15','Metric Ton','Reservasi','ShipTo bentrok dengan Plumpang'
	if perihal == 'L15':
		tripBaru = 'Ada pengisian Liter 15(L15) dengan data sebagai berikut: '
		docPath = r'D:\test\NOTA.docx'
		doc = docx.Document(docPath)
		doc.tables[0].cell(1,4).text = 'L15'
		doc.paragraphs[0].runs[3].text = noNota
		doc.paragraphs[7].text = tripBaru
		#row1
		doc.tables[0].cell(2,0).text = nomer1
		doc.tables[0].cell(2,1).text = mobil1
		doc.tables[0].cell(2,2).text = nomerLO1
		doc.tables[0].cell(2,3).text = produkLO1
		doc.tables[0].cell(2,4).text = ml1
		doc.tables[0].cell(2,5).text = volume1
		doc.tables[0].cell(2,6).text = keterangan
		#row2
		doc.tables[0].cell(3,0).text = nomer2
		doc.tables[0].cell(3,1).text = mobil2
		doc.tables[0].cell(3,2).text = nomerLO2
		doc.tables[0].cell(3,3).text = produkLO2
		doc.tables[0].cell(3,4).text = ml2
		doc.tables[0].cell(3,5).text = volume2
		doc.tables[0].cell(3,6).text = keterangan2
		#row3
		doc.tables[0].cell(4,0).text = nomer3
		doc.tables[0].cell(4,1).text = mobil3
		doc.tables[0].cell(4,2).text = nomerLO3
		doc.tables[0].cell(4,3).text = produkLO3
		doc.tables[0].cell(4,4).text = ml3
		doc.tables[0].cell(4,5).text = volume3
		doc.tables[0].cell(4,6).text = keterangan3
		#row4
		doc.tables[0].cell(5,0).text = nomer4
		doc.tables[0].cell(5,1).text = mobil4
		doc.tables[0].cell(5,2).text = nomerLO4
		doc.tables[0].cell(5,3).text = produkLO4
		doc.tables[0].cell(5,4).text = ml4
		doc.tables[0].cell(5,5).text = volume4
		doc.tables[0].cell(5,6).text = keterangan4
		#row5
		doc.tables[0].cell(6,0).text = nomer5
		doc.tables[0].cell(6,1).text = mobil5
		doc.tables[0].cell(6,2).text = nomerLO5
		doc.tables[0].cell(6,3).text = produkLO5
		doc.tables[0].cell(6,4).text = ml5
		doc.tables[0].cell(6,5).text = volume5
		doc.tables[0].cell(6,6).text = keterangan5
		#row6
		doc.tables[0].cell(7,0).text = nomer6
		doc.tables[0].cell(7,1).text = mobil6
		doc.tables[0].cell(7,2).text = nomerLO6
		doc.tables[0].cell(7,3).text = produkLO6
		doc.tables[0].cell(7,4).text = ml6
		doc.tables[0].cell(7,5).text = volume6
		doc.tables[0].cell(7,6).text = keterangan6
		#row7
		doc.tables[0].cell(8,0).text = nomer7
		doc.tables[0].cell(8,1).text = mobil7
		doc.tables[0].cell(8,2).text = nomerLO7
		doc.tables[0].cell(8,3).text = produkLO7
		doc.tables[0].cell(8,4).text = ml7
		doc.tables[0].cell(8,5).text = volume7
		doc.tables[0].cell(8,6).text = keterangan7
		#4.interpolasi tempat dan tanggal
		tempatTanggal = 'Jakarta, {0}'.format(today.strftime('%d %B %Y'))
		doc.paragraphs[11].runs[0].text = tempatTanggal


	#2.2 Metric ton[t]
	elif perihal == 'Metric Ton':
		mtAsu = 'Ada pengisian Metric Ton dengan data sebagai berikut:'
		docPath = r'D:\test\NOTA.docx'
		doc = docx.Document(docPath)
		doc.paragraphs[0].runs[3].text = noNota
		doc.tables[0].cell(1,4).text = 'MetricTon[t]'
		doc.paragraphs[7].text = mtAsu
		#row1
		doc.tables[0].cell(2,0).text = nomer1
		doc.tables[0].cell(2,1).text = mobil1
		doc.tables[0].cell(2,2).text = nomerLO1
		doc.tables[0].cell(2,3).text = produkLO1
		doc.tables[0].cell(2,4).text = ml1
		doc.tables[0].cell(2,5).text = volume1
		doc.tables[0].cell(2,6).text = keterangan
		#row2
		doc.tables[0].cell(3,0).text = nomer2
		doc.tables[0].cell(3,1).text = mobil2
		doc.tables[0].cell(3,2).text = nomerLO2
		doc.tables[0].cell(3,3).text = produkLO2
		doc.tables[0].cell(3,4).text = ml2
		doc.tables[0].cell(3,5).text = volume2
		doc.tables[0].cell(3,6).text = keterangan2
		#row3
		doc.tables[0].cell(4,0).text = nomer3
		doc.tables[0].cell(4,1).text = mobil3
		doc.tables[0].cell(4,2).text = nomerLO3
		doc.tables[0].cell(4,3).text = produkLO3
		doc.tables[0].cell(4,4).text = ml3
		doc.tables[0].cell(4,5).text = volume3
		doc.tables[0].cell(4,6).text = keterangan3
		#row4
		doc.tables[0].cell(5,0).text = nomer4
		doc.tables[0].cell(5,1).text = mobil4
		doc.tables[0].cell(5,2).text = nomerLO4
		doc.tables[0].cell(5,3).text = produkLO4
		doc.tables[0].cell(5,4).text = ml4
		doc.tables[0].cell(5,5).text = volume4
		doc.tables[0].cell(5,6).text = keterangan4
		#row5
		doc.tables[0].cell(6,0).text = nomer5
		doc.tables[0].cell(6,1).text = mobil5
		doc.tables[0].cell(6,2).text = nomerLO5
		doc.tables[0].cell(6,3).text = produkLO5
		doc.tables[0].cell(6,4).text = ml5
		doc.tables[0].cell(6,5).text = volume5
		doc.tables[0].cell(6,6).text = keterangan5
		#row6
		doc.tables[0].cell(7,0).text = nomer6
		doc.tables[0].cell(7,1).text = mobil6
		doc.tables[0].cell(7,2).text = nomerLO6
		doc.tables[0].cell(7,3).text = produkLO6
		doc.tables[0].cell(7,4).text = ml6
		doc.tables[0].cell(7,5).text = volume6
		doc.tables[0].cell(7,6).text = keterangan6
		#row7
		doc.tables[0].cell(8,0).text = nomer7
		doc.tables[0].cell(8,1).text = mobil7
		doc.tables[0].cell(8,2).text = nomerLO7
		doc.tables[0].cell(8,3).text = produkLO7
		doc.tables[0].cell(8,4).text = ml7
		doc.tables[0].cell(8,5).text = volume7
		doc.tables[0].cell(8,6).text = keterangan7
		#4.interpolasi tempat dan tanggal
		tempatTanggal = 'Jakarta, {0}'.format(today.strftime('%d %B %Y'))
		doc.paragraphs[11].runs[0].text = tempatTanggal


	#2.3 Reservasi
	elif perihal == 'Reservasi':
		doPe = 'Ada pengisian reservasi dengan data sebagai berikut:'
		docPath = r'D:\test\NOTA2.docx'
		doc = docx.Document(docPath)
		doc.paragraphs[0].runs[3].text = noNota
		doc.paragraphs[7].text = doPe
		#row1
		doc.tables[0].cell(1,0).text = nomer1
		doc.tables[0].cell(1,1).text = mobil1
		doc.tables[0].cell(1,2).text = nomerLO1
		doc.tables[0].cell(1,3).text = produkLO1
		doc.tables[0].cell(1,4).text = volume1
		doc.tables[0].cell(1,5).text = keterangan
		#row2
		doc.tables[0].cell(2,0).text = nomer2
		doc.tables[0].cell(2,1).text = mobil2
		doc.tables[0].cell(2,2).text = nomerLO2
		doc.tables[0].cell(2,3).text = produkLO2
		doc.tables[0].cell(2,4).text = volume2
		doc.tables[0].cell(2,5).text = keterangan2
		#row3
		doc.tables[0].cell(3,0).text = nomer3
		doc.tables[0].cell(3,1).text = mobil3
		doc.tables[0].cell(3,2).text = nomerLO3
		doc.tables[0].cell(3,3).text = produkLO3
		doc.tables[0].cell(3,4).text = volume3
		doc.tables[0].cell(3,5).text = keterangan3
		#row4
		doc.tables[0].cell(4,0).text = nomer4
		doc.tables[0].cell(4,1).text = mobil4
		doc.tables[0].cell(4,2).text = nomerLO4
		doc.tables[0].cell(4,3).text = produkLO4
		doc.tables[0].cell(4,4).text = volume4
		doc.tables[0].cell(4,5).text = keterangan4
		#row5
		doc.tables[0].cell(5,0).text = nomer5
		doc.tables[0].cell(5,1).text = mobil5
		doc.tables[0].cell(5,2).text = nomerLO5
		doc.tables[0].cell(5,3).text = produkLO5
		doc.tables[0].cell(5,4).text = volume5
		doc.tables[0].cell(5,5).text = keterangan5
		#row6
		doc.tables[0].cell(6,0).text = nomer6
		doc.tables[0].cell(6,1).text = mobil6
		doc.tables[0].cell(6,2).text = nomerLO6
		doc.tables[0].cell(6,3).text = produkLO6
		doc.tables[0].cell(6,4).text = volume6
		doc.tables[0].cell(6,5).text = keterangan6
		#row7
		#doc.tables[0].cell(7,0).text = nomer7
		#doc.tables[0].cell(7,1).text = mobil7
		#doc.tables[0].cell(7,2).text = nomerLO7
		#doc.tables[0].cell(7,3).text = produkLO7
		#doc.tables[0].cell(7,4).text = volume7
		#4.interpolasi tempat dan tanggal
		tempatTanggal = 'Jakarta, {0}'.format(today.strftime('%d %B %Y'))
		doc.paragraphs[10].runs[0].text = tempatTanggal
		

	#2.4 ShipTo bentrok
	elif perihal == 'ShipTo bentrok dengan Plumpang':
		konservasi = 'Terdapat ship to bentrok antara carrier Instalasi Jakarta Group dan Instalasi Tanjung Priok dengan data sebagai berikut:'
		docPath = r'D:\test\NOTA2.docx'
		doc = docx.Document(docPath)
		doc.paragraphs[0].runs[3].text = noNota
		doc.paragraphs[7].text = konservasi
		#row1
		doc.tables[0].cell(1,0).text = nomer1
		doc.tables[0].cell(1,1).text = mobil1
		doc.tables[0].cell(1,2).text = nomerLO1
		doc.tables[0].cell(1,3).text = produkLO1
		doc.tables[0].cell(1,4).text = volume1
		doc.tables[0].cell(1,5).text = keterangan
		#row2
		doc.tables[0].cell(2,0).text = nomer2
		doc.tables[0].cell(2,1).text = mobil2
		doc.tables[0].cell(2,2).text = nomerLO2
		doc.tables[0].cell(2,3).text = produkLO2
		doc.tables[0].cell(2,4).text = volume2
		doc.tables[0].cell(2,5).text = keterangan2
		#row3
		doc.tables[0].cell(3,0).text = nomer3
		doc.tables[0].cell(3,1).text = mobil3
		doc.tables[0].cell(3,2).text = nomerLO3
		doc.tables[0].cell(3,3).text = produkLO3
		doc.tables[0].cell(3,4).text = volume3
		doc.tables[0].cell(3,5).text = keterangan3
		#row4
		doc.tables[0].cell(4,0).text = nomer4
		doc.tables[0].cell(4,1).text = mobil4
		doc.tables[0].cell(4,2).text = nomerLO4
		doc.tables[0].cell(4,3).text = produkLO4
		doc.tables[0].cell(4,4).text = volume4
		doc.tables[0].cell(4,5).text = keterangan4
		#row5
		doc.tables[0].cell(5,0).text = nomer5
		doc.tables[0].cell(5,1).text = mobil5
		doc.tables[0].cell(5,2).text = nomerLO5
		doc.tables[0].cell(5,3).text = produkLO5
		doc.tables[0].cell(5,4).text = volume5
		doc.tables[0].cell(5,5).text = keterangan5
		#row6
		doc.tables[0].cell(6,0).text = nomer6
		doc.tables[0].cell(6,1).text = mobil6
		doc.tables[0].cell(6,2).text = nomerLO6
		doc.tables[0].cell(6,3).text = produkLO6
		doc.tables[0].cell(6,4).text = volume6
		doc.tables[0].cell(6,5).text = keterangan6
		#row7
		#doc.tables[0].cell(7,0).text = nomer7
		#doc.tables[0].cell(7,1).text = mobil7
		#doc.tables[0].cell(7,2).text = nomerLO7
		#doc.tables[0].cell(7,3).text = produkLO7
		#doc.tables[0].cell(7,4).text = volume7
		#4.interpolasi tempat dan tanggal
		tempatTanggal = 'Jakarta, {0}'.format(today.strftime('%d %B %Y'))
		doc.paragraphs[10].runs[0].text = tempatTanggal

    #makeStructuredFolder
	yyPath = r'D:\NOTA TRIP NUMBER BARU\%s'% thisYear
	if not os.path.exists(yyPath):
		os.makedirs(yyPath)
	theMonth = today.strftime('%B')	
	mmPath = r'D:\NOTA TRIP NUMBER BARU\%s\%s' % (thisYear,theMonth)
	
	if not os.path.exists(mmPath):
		os.makedirs(mmPath)

	theDay = today.strftime('%d')
	ddPath = r'D:\NOTA TRIP NUMBER BARU\%s\%s\%s' % (thisYear,theMonth,theDay)
	if not os.path.exists(ddPath):
		os.makedirs(ddPath)
	docxSavedFilePath = r'D:\NOTA TRIP NUMBER BARU\%s\%s\%s\Nota %s.docx' % (thisYear,theMonth,theDay,nomerNota)
	doc.save(docxSavedFilePath)
	
        #convert to pdf
	
	pdfSavedFilePath = r'D:\NOTA TRIP NUMBER BARU\%s\%s\%s\Nota %s.pdf'% (thisYear,theMonth,theDay,nomerNota)
	wdFormatPDF = 17
	wordy = comtypes.client.CreateObject('Word.Application',dynamic=True)
	wordy.Visible = True
	in_file = os.path.abspath(docxSavedFilePath)
	out_file = os.path.abspath(pdfSavedFilePath)
	wordy.Documents.Open(in_file)
	wordy.Documents[0].SaveAs(out_file,17)
	wordy.Documents[0].Close()
	os.remove(docxSavedFilePath)
	#email stuff
	fromAddr = 'plumpang.automail.com@gmail.com'
	toAddr = 'evi@cpan.biz'
	testAddr = 'vansandriawan@gmail.com'
	pswd = 'prisela99'
	namaFile = 'NOTA %s PMB.pdf' % nomerNota
	msg = MIMEMultipart()
	msg['Subject']= 'NOTA %s PMB'% nomerNota
	#pdfAttachment
	fp = open(pdfSavedFilePath,'rb')
	att = MIMEApplication(fp.read(),_subtype='pdf')
	fp.close()
	att.add_header('Content-Disposition','attachment',filename=namaFile)
	msg.attach(att)
	#sending the shit
	mailer = smtplib.SMTP('smtp.gmail.com',587)
	mailer.starttls()
	mailer.ehlo()
	mailer.login(fromAddr,pswd)
	mailer.sendmail(fromAddr,toAddr,msg.as_string())
	mailer.quit()
	#deleteEntry
	notaEntry.delete(0,END)
	noEntry1.delete(0,END)
	nopolEntry1.delete(0,END)
	noloEntry1.delete(0,END)
	prodEntry1.delete(0,END)
	mlEntry1.delete(0,END)
	volEntry1.delete(0,END)
	ketEntry.delete(0,END)
	#row5
	noEntry2.delete(0,END)
	nopolEntry2.delete(0,END)
	noloEntry2.delete(0,END)
	prodEntry2.delete(0,END)
	mlEntry2.delete(0,END)
	volEntry2.delete(0,END)
	ketEntry2.delete(0,END)
	#row6
	noEntry3.delete(0,END)
	nopolEntry3.delete(0,END)
	noloEntry3.delete(0,END)
	prodEntry3.delete(0,END)
	mlEntry3.delete(0,END)
	volEntry3.delete(0,END)
	ketEntry3.delete(0,END)
	#row7
	noEntry4.delete(0,END)
	nopolEntry4.delete(0,END)
	noloEntry4.delete(0,END)
	prodEntry4.delete(0,END)
	mlEntry4.delete(0,END)
	volEntry4.delete(0,END)
	ketEntry4.delete(0,END)
	#row8
	noEntry5.delete(0,END)
	nopolEntry5.delete(0,END)
	noloEntry5.delete(0,END)
	prodEntry5.delete(0,END)
	mlEntry5.delete(0,END)
	volEntry5.delete(0,END)
	ketEntry5.delete(0,END)
	#row9
	noEntry6.delete(0,END)
	nopolEntry6.delete(0,END)
	noloEntry6.delete(0,END)
	prodEntry6.delete(0,END)
	mlEntry6.delete(0,END)
	volEntry6.delete(0,END)
	ketEntry6.delete(0,END)
	#row10
	noEntry7.delete(0,END)
	nopolEntry7.delete(0,END)
	noloEntry7.delete(0,END)
	prodEntry7.delete(0,END)
	mlEntry7.delete(0,END)
	volEntry7.delete(0,END)
	ketEntry7.delete(0,END)
	t2=Toplevel(tab1)
	Label(t2,text='Email has been sent\n Please Chill the F out').pack(padx=50,pady=50)
	t2.transient(tab1)
Button(tab1,text='generate',command=runShit).grid(row=11, column=3)
tab1.mainloop()
#///
