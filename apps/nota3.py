import sys
sys.path.insert(0,'../Date')
from structuredFolder import rome,structFolderMan
from allDateFormat import getNow,getDashDate
from emailThingy import mailAttachment,logIn,testAddr,eAddress
import docx
from docx.shared import Inches
import comtypes.client
import os
from email.mime.multipart import MIMEMultipart
def takeVal(valueList):
    d = []
    for i in valueList:
        if i.winfo_class() == 'Entry':
            d.append(str(i.get()))
    return d
def docxToPdf(docxPath,docxFolder,nomerNota):
    # convert to pdf
    pdfSavedFilePath = docxFolder+'\\Nota %s.pdf' % (nomerNota)
    #wdFormatPDF = 17
    wordy = comtypes.client.CreateObject('Word.Application', dynamic=True)
    wordy.Visible = False #True
    in_file = os.path.abspath(docxPath)
    out_file = os.path.abspath(pdfSavedFilePath)
    wordy.Documents.Open(in_file)
    wordy.Documents[0].SaveAs(out_file, 17)
    wordy.Documents[0].Close()
    os.remove(docxPath)
    return pdfSavedFilePath
def mailContentNota(pdfPath,nomerNota,site):
    namaFile = 'NOTA {0} {1}'.format(nomerNota,site)
    msg = MIMEMultipart()
    msg['Subject'] = namaFile
    stageName = 'NOTA {0} {1}.pdf'.format(nomerNota,site)
    att1 = mailAttachment(pdfPath, stageName)
    msg.attach(att1)
    return msg
def toDocx(nomerNota,noNota,perihal,tableList,tempatTanggal,imageList,site):
    docPath = "..\..\\template\\nota\\NOTA.docx"
    doc = docx.Document(docPath)
    doc.paragraphs[0].runs[3].text = noNota #
    doc.paragraphs[7].text = perihal #
    # 3.interpolasi Tabel
    for j in range(6):
        doc.tables[0].cell(1, j).text = tableList[j]
        doc.tables[0].cell(2, j).text = tableList[j+6]
        doc.tables[0].cell(3, j).text = tableList[j+12]
        doc.tables[0].cell(4, j).text = tableList[j+18]
        doc.tables[0].cell(5, j).text = tableList[j+24]
        doc.tables[0].cell(6, j).text = tableList[j+30]
        doc.tables[0].cell(7, j).text = tableList[j+36]
    doc.paragraphs[11].runs[0].text = tempatTanggal
    for i in imageList:
        addPar = doc.add_paragraph() #from this
        addRun = addPar.add_run()
        addRun.add_picture(i,width=Inches(6.5)) #to this
    docxFolder = structFolderMan('NOTA TRIP NUMBER BARU',site,getDashDate())
    docxSavedFilePath = docxFolder+'\\Nota %s.docx'%nomerNota
    doc.save(docxSavedFilePath)
    pdFile = docxToPdf(docxSavedFilePath, docxFolder, nomerNota)
    return pdFile
def siteProfile(site):
	if site == 'PLP':
		namaNota = 'PLP-MNJ'
		kota = 'Jakarta'
		return namaNota,kota
	elif site == 'BLG':
		namaNota = 'BLG-MNJ'#not fixed
		kota = 'Indramayu'
		return namaNota,kota
	elif site == 'TGR':
		namaNota = 'TGR-MNJ'#not fixed
		kota = 'Merak'
		return namaNota,kota
	elif site == 'UJB':
		namaNota = 'UJB-MNJ'#not fixed
		kota = 'Bandung'
		return namaNota,kota
	elif site == 'SBY':
		namaNota = 'SBY-MNJ'#not fixed
		kota = 'Surabaya'
		return namaNota,kota
	elif site == 'BYL':
		namaNota = 'BYL-MNJ'#not fixed
		kota = 'Boyolali'
		return namaNota,kota
	elif site == 'PMB':
		namaNota = 'PMB-MNJ'#not fixed
		kota = 'Jakarta'
		return namaNota,kota
	elif site == 'MDN':
		namaNota = 'MDN-MNJ'#not fixed
		kota = 'Medan'
		return namaNota,kota
	elif site == 'KTP':
		namaNota = 'KTP-MNJ'#not fixed
		kota = 'Palembang'
		return namaNota,kota
	elif site == 'PJG':
		namaNota = 'PJG-MNJ'#not fixed
		kota = 'Lampung'
		return namaNota,kota
def nota(tableList,opt_perihal,imageList):
    today = getNow()
    # 1.interpolasi nomor nota
    thisMonth = today.strftime('%m')
    romanizing = rome(thisMonth)
    notaYear = today.strftime('%Y')
    nomerNota = tableList[0]
    noNota = '{0}/{3}/{1}/{2}'.format(nomerNota, romanizing, notaYear, namaNota)
    tempatTanggal = '{1}, {0}'.format(today.strftime('%d %B %Y'),kota)
    del tableList[0]
    # 2.interpolasi Perihal (conditional)
    # 2.1 trip number baru
    if opt_perihal == 'Trip Number tanpa LO':
        perihal = 'Ada pengisian yang menyebabkan Temperature is inconsistent with given Temperature in Metered Qty sehingga harus dibuatkan trip number baru dengan data sebagai berikut : '
        #doc.paragraphs[7].text = tripBaru
        pdFile = toDocx(nomerNota,noNota,perihal,tableList,tempatTanggal,imageList)
        fromAddr, toAddr, pswd = eAddress('PLP_NOTA')
        msg = mailContentNota(pdFile, nomerNota)
        logIn(fromAddr, toAddr, pswd, msg)
    # 2.2 MT ASU
    elif opt_perihal == 'MT UNIK':
        mobilAsu = tableList[1]
        perihal = 'Terdapat pengisian pada MT {0} (MT Unik)  yang menyebabkan schedule tersebut tidak dapat masuk secara otomatis ke Omega:'.format(mobilAsu)
        #doc.paragraphs[7].text = mtAsu
        pdFile = toDocx(nomerNota, noNota, perihal, tableList, tempatTanggal,imageList)
        fromAddr, toAddr, pswd = eAddress('PLP_NOTA')#test
        msg = mailContentNota(pdFile, nomerNota)
        logIn(fromAddr, toAddr, pswd, msg)
    # 2.3 DO pecah di bawah kapasitas
    elif opt_perihal == 'DO Pecah':
        noMT = tableList[1]
        perihal = 'Terdapat pengisian do pecah dibawah kapasitas pada MT {0}, yang menyebabkan schedule tersebut tidak dapat masuk secara otomatis ke Omega :'.format(noMT)
        #doc.paragraphs[7].text = doPe
        pdFile = toDocx(nomerNota, noNota, perihal, tableList, tempatTanggal,imageList)
        fromAddr, toAddr, pswd = eAddress('PLP_NOTA')
        msg = mailContentNota(pdFile, nomerNota)
        logIn(fromAddr, toAddr, pswd, msg)
    # 2.4 Konsinyasi dan Reservasi
    elif opt_perihal == 'Konservasi':
        perihal = 'Terdapat pengisian konsinyasi/reservasi yang  menyebabkan schedule tersebut tidak dapat masuk secara otomatis ke MPV :'
        #doc.paragraphs[7].text = konservasi
        pdFile = toDocx(nomerNota, noNota, perihal, tableList, tempatTanggal,imageList)
        fromAddr, toAddr, pswd = eAddress('PLP_NOTA')
        msg = mailContentNota(pdFile, nomerNota)
        logIn(fromAddr, toAddr, pswd, msg)
    # 3.interpolasi Tabel
    #for j in range(6):
     #   doc.tables[0].cell(1, j).text = tableList[j]
      #  doc.tables[0].cell(2, j).text = tableList[j+6]
       # doc.tables[0].cell(3, j).text = tableList[j+12]
        #doc.tables[0].cell(4, j).text = tableList[j+12]
        #doc.tables[0].cell(5, j).text = tableList[j+18]
        #doc.tables[0].cell(6, j).text = tableList[j+24]
        #doc.tables[0].cell(7, j).text = tableList[j+30]
    # email stuff

    # pdfAttachment
    #fp = open(pdfSavedFilePath, 'rb')
    #att = MIMEApplication(fp.read(), _subtype='pdf')
    #fp.close()
    #att.add_header('Content-Disposition', 'attachment', filename=namaFile)
    #msg.attach(att)
